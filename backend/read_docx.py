import docx, sys

doc = docx.Document(r"c:\Projects\scheduleSYS\База данных\Кураторы и контакты Инф.технол. 2026.docx")
for i, t in enumerate(doc.tables):
    print(f"TABLE {i}")
    for j, r in enumerate(t.rows):
        print("|".join(c.text.strip() for c in r.cells))
print("---PARAGRAPHS---")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text.strip())
