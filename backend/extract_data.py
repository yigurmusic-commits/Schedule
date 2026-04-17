"""
Скрипт для извлечения данных из файлов папки 'База данных'.
Сохраняет результат в output_data.txt
"""
import os
import sys

from docx import Document
import xlrd
import openpyxl

BASE_DIR = os.path.join(os.path.dirname(__file__), "..", "База данных")
OUTPUT_FILE = os.path.join(os.path.dirname(__file__), "output_data.txt")


def read_docx_tables(filepath, f):
    doc = Document(filepath)
    f.write(f"\n{'='*80}\n")
    f.write(f"ФАЙЛ: {os.path.basename(filepath)}\n")
    f.write(f"Параграфов: {len(doc.paragraphs)}, Таблиц: {len(doc.tables)}\n")
    
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            f.write(f"  П{i}: {p.text.strip()[:150]}\n")
    
    for ti, table in enumerate(doc.tables):
        f.write(f"\n  --- Таблица {ti+1} ({len(table.rows)} строк x {len(table.columns)} столбцов) ---\n")
        for ri, row in enumerate(table.rows):
            if ri > 30:
                f.write(f"  ... ещё {len(table.rows) - 30} строк\n")
                break
            cells = [cell.text.strip().replace('\n', ' ')[:50] for cell in row.cells]
            f.write(f"  R{ri}: {' | '.join(cells)}\n")


def read_xls(filepath, f):
    f.write(f"\n{'='*80}\n")
    f.write(f"ФАЙЛ: {os.path.basename(filepath)}\n")
    wb = xlrd.open_workbook(filepath)
    for sheet in wb.sheets():
        f.write(f"\n  --- Лист: '{sheet.name}' ({sheet.nrows} строк x {sheet.ncols} столбцов) ---\n")
        for ri in range(min(sheet.nrows, 30)):
            cells = []
            for ci in range(min(sheet.ncols, 15)):
                val = sheet.cell_value(ri, ci)
                if val:
                    cells.append(str(val)[:50])
                else:
                    cells.append("")
            f.write(f"  R{ri}: {' | '.join(cells)}\n")
        if sheet.nrows > 30:
            f.write(f"  ... ещё {sheet.nrows - 30} строк\n")


def read_xlsx(filepath, f):
    f.write(f"\n{'='*80}\n")
    f.write(f"ФАЙЛ: {os.path.basename(filepath)}\n")
    wb = openpyxl.load_workbook(filepath, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        f.write(f"\n  --- Лист: '{sheet_name}' ({ws.max_row} строк x {ws.max_column} столбцов) ---\n")
        for ri, row in enumerate(ws.iter_rows(max_row=min(ws.max_row, 30), values_only=True), 1):
            cells = [str(c)[:50] if c is not None else "" for c in row[:15]]
            f.write(f"  R{ri}: {' | '.join(cells)}\n")
        if ws.max_row > 30:
            f.write(f"  ... ещё {ws.max_row - 30} строк\n")


if __name__ == "__main__":
    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        for fname in sorted(os.listdir(BASE_DIR)):
            fpath = os.path.join(BASE_DIR, fname)
            if not os.path.isfile(fpath):
                continue
            
            ext = os.path.splitext(fname)[1].lower()
            try:
                if ext == ".docx":
                    read_docx_tables(fpath, f)
                elif ext == ".xls":
                    read_xls(fpath, f)
                elif ext == ".xlsx":
                    read_xlsx(fpath, f)
            except Exception as e:
                f.write(f"\nОШИБКА при чтении {fname}: {e}\n")
    
    print(f"Готово! Результат сохранён в {OUTPUT_FILE}")
